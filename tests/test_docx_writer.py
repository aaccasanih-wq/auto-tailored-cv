"""Tests for src/render/docx_writer.py — using a real built docx fixture."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.profile.cv_reader import read_cv
from src.render.docx_writer import write_tailored_docx, _set_paragraph_text, _set_cell_text


@pytest.fixture
def base_docx(tmp_path: Path) -> Path:
    """Build a CV-docx fixture mirroring the real one's structure."""
    doc = Document()
    doc.add_paragraph("AXEL AARON CCASANI HUACHUA")
    doc.add_paragraph("986 531 450 | aaronaxel810@gmail.com")
    doc.add_paragraph("En búsqueda de un puesto en Data Science")
    doc.add_paragraph("")
    doc.add_paragraph("EDUCACIÓN")
    doc.add_paragraph("Lic. en Economía | 2021 – 2026")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Universidad del Pacífico"
    t.rows[0].cells[1].text = "2021 – 2026"

    doc.add_paragraph("EXPERIENCIA LABORAL")
    doc.add_paragraph("Automaticé los procesos de validación de facturas.")
    doc.add_paragraph("Extraje y analicé datos de SAP.")
    t2 = doc.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "Metso Perú — Practicante"
    t2.rows[0].cells[1].text = "Nov 2024 – Feb 2025"

    doc.add_paragraph("HABILIDADES & HERRAMIENTAS")
    t3 = doc.add_table(rows=2, cols=2)
    t3.rows[0].cells[0].text = "Python"
    t3.rows[0].cells[1].text = "Pandas, NumPy, Jupyter"
    t3.rows[1].cells[0].text = "Excel"
    t3.rows[1].cells[1].text = "SAP"

    path = tmp_path / "base_cv.docx"
    doc.save(str(path))
    return path


class TestLowLevel:
    def test_set_paragraph_text_preserves_run_count(self, tmp_path):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("hello")
        p.add_run(" ")
        p.add_run("world")
        path = tmp_path / "tmp.docx"
        doc.save(str(path))

        _set_paragraph_text(p, "new text")
        assert p.text == "new text"
        # First run keeps formatting, others are blanked but not removed
        texts = [r.text for r in p.runs]
        assert texts == ["new text", "", ""]

    def test_set_paragraph_text_no_runs(self, tmp_path):
        doc = Document()
        p = doc.add_paragraph()
        # No runs yet
        _set_paragraph_text(p, "added")
        assert p.text == "added"


class TestWriteTailoredDocx:
    def test_summary_rewritten(self, base_docx: Path, tmp_path: Path):
        out = tmp_path / "tailored.docx"
        tailored = {
            "summary": "En búsqueda de un puesto en Data Engineering",
            "sections": [],
        }
        write_tailored_docx(base_docx, tailored, out)
        profile = read_cv(out)
        assert profile.summary == "En búsqueda de un puesto en Data Engineering"

    def test_name_unchanged(self, base_docx: Path, tmp_path: Path):
        out = tmp_path / "tailored.docx"
        tailored = {"summary": "", "sections": []}
        write_tailored_docx(base_docx, tailored, out)
        assert read_cv(out).name == "AXEL AARON CCASANI HUACHUA"

    def test_contact_unchanged(self, base_docx: Path, tmp_path: Path):
        out = tmp_path / "tailored.docx"
        tailored = {"summary": "", "sections": []}
        write_tailored_docx(base_docx, tailored, out)
        assert "aaronaxel810@gmail.com" in read_cv(out).contact

    def test_section_paragraph_rewritten(self, base_docx: Path, tmp_path: Path):
        out = tmp_path / "tailored.docx"
        tailored = {
            "summary": "",
            "sections": [
                {"title": "EDUCACIÓN", "paragraphs": ["Bachelor in Economics | 2021 – 2026"],
                 "tables": [[["", ""]]]},
                {"title": "EXPERIENCIA LABORAL",
                 "paragraphs": ["Automated invoice validation workflow.",
                                "Extracted data from SAP."],
                 "tables": [[["", ""]]]},
                {"title": "HABILIDADES & HERRAMIENTAS", "paragraphs": [],
                 "tables": [[["", ""], ["", ""]]]},
            ],
        }
        write_tailored_docx(base_docx, tailored, out)
        profile = read_cv(out)
        exp = next(s for s in profile.sections if s.title == "EXPERIENCIA LABORAL")
        assert exp.paragraphs[0] == "Automated invoice validation workflow."
        assert exp.paragraphs[1] == "Extracted data from SAP."

    def test_section_table_replaced(self, base_docx: Path, tmp_path: Path):
        out = tmp_path / "tailored.docx"
        new_skills = ["Python", "Pandas, NumPy, Streamlit, Selenium", "Excel", "SAP"]
        tailored = {
            "summary": "",
            "sections": [
                {"title": "EDUCACIÓN", "paragraphs": ["x"], "tables": [[["a", "b"]]]},
                {"title": "EXPERIENCIA LABORAL", "paragraphs": ["x", "y"], "tables": [[["c", "d"]]]},
                {"title": "HABILIDADES & HERRAMIENTAS", "paragraphs": [],
                 "tables": [[[new_skills[0], new_skills[1]], [new_skills[2], new_skills[3]]]]},
            ],
        }
        write_tailored_docx(base_docx, tailored, out)
        profile = read_cv(out)
        hab = next(s for s in profile.sections if s.title == "HABILIDADES & HERRAMIENTAS")
        assert hab.tables[0][0] == ["Python", "Pandas, NumPy, Streamlit, Selenium"]
        assert hab.tables[0][1] == ["Excel", "SAP"]

    def test_cell_with_multiple_paragraphs_collapsed(self, base_docx: Path, tmp_path: Path):
        # The real CV's skills cell has a multi-line text. Build tailored text
        # as a single string and confirm cell.text matches exactly after write.
        out = tmp_path / "tailored.docx"
        big_skill_value = (
            "Python (Pandas, NumPy, Matplotlib, Seaborn, Selenium, APIs/Requests, "
            "GeoPandas, Folium, Raster, Langchain, etc), SQL, R (RStudio)"
        )
        tailored = {
            "summary": "",
            "sections": [
                {"title": "EDUCACIÓN", "paragraphs": ["x"], "tables": [[["a", "b"]]]},
                {"title": "EXPERIENCIA LABORAL", "paragraphs": ["x", "y"], "tables": [[["c", "d"]]]},
                {"title": "HABILIDADES & HERRAMIENTAS", "paragraphs": [],
                 "tables": [[["Programming", big_skill_value]]]},
            ],
        }
        # Build base docx with a multi-paragraph cell
        doc = Document(str(base_docx))
        hab_table = doc.tables[-1]
        # Reset the cell to multi-paragraph text by hand:
        cell = hab_table.rows[0].cells[1]
        # cell currently has one paragraph "Pandas, NumPy, Jupyter"
        cell.paragraphs[0].text = ""  # clear via setting via run not preserved; ok for test
        cell.add_paragraph("More text line 2")
        new_base = tmp_path / "base_multi_para.docx"
        doc.save(str(new_base))

        write_tailored_docx(new_base, tailored, out)
        out_doc = Document(str(out))
        cells = out_doc.tables[-1].rows[0].cells
        # cell.text concatenates paragraphs with \n; ours should be collapsed to one
        # so cell text == big_skill_value exactly (one paragraph).
        assert cells[1].text == big_skill_value

    def test_preserves_section_titles_unchanged(self, base_docx: Path, tmp_path: Path):
        out = tmp_path / "tailored.docx"
        tailored = {
            "summary": "", "sections": [
                {"title": "EDUCACIÓN", "paragraphs": ["x"], "tables": [[["a", "b"]]]},
                {"title": "EXPERIENCIA LABORAL", "paragraphs": ["x", "y"], "tables": [[["c", "d"]]]},
                {"title": "HABILIDADES & HERRAMIENTAS", "paragraphs": [], "tables": [[["", ""], ["", ""]]]},
            ],
        }
        write_tailored_docx(base_docx, tailored, out)
        titles = [s.title for s in read_cv(out).sections]
        assert titles == ["EDUCACIÓN", "EXPERIENCIA LABORAL", "HABILIDADES & HERRAMIENTAS"]

    def test_original_file_not_modified(self, base_docx: Path, tmp_path: Path):
        original_summary = read_cv(base_docx).summary
        out = tmp_path / "tailored.docx"
        tailored = {"summary": "totally new summary", "sections": []}
        write_tailored_docx(base_docx, tailored, out)
        # The base docx fixture should be unchanged on disk
        assert read_cv(base_docx).summary == original_summary