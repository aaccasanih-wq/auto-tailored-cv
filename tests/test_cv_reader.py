"""Tests for src/profile/cv_reader.py"""

from pathlib import Path

import pytest
from docx import Document

from src.profile.cv_reader import (
    CVProfile,
    CVSection,
    _is_heading,
    iter_block_items,
    read_cv,
)


class TestIsHeading:
    def test_english_uppercase(self):
        assert _is_heading("EDUCACIÓN") is True

    def test_uppercase_with_spaces(self):
        assert _is_heading("EXPERIENCIA LABORAL") is True

    def test_uppercase_with_ampersand(self):
        assert _is_heading("HABILIDADES & HERRAMIENTAS") is True

    def test_lowercase_mixed(self):
        assert _is_heading("Rastreador de Gastos Automatizado con IA") is False

    def test_lowercase_accented(self):
        assert _is_heading("Ingeniero de Datos Sénior") is False

    def test_empty(self):
        assert _is_heading("") is False

    def test_only_digits(self):
        assert _is_heading("2021 – 2026") is False

    def test_only_punctuation(self):
        assert _is_heading("---") is False


@pytest.fixture
def tmp_cv(tmp_path: Path) -> Path:
    """Build a small CV .docx in tmp_path mirroring the real CV convention."""
    doc = Document()
    doc.add_paragraph("ALEX SAMPLE CANDIDATE")
    doc.add_paragraph("555 0100 | candidate@example.com")
    doc.add_paragraph("En búsqueda de un puesto en Data Science")

    doc.add_paragraph("EDUCACIÓN")
    doc.add_paragraph("Lic. en Economía | 2021 – 2026")
    table_edu = doc.add_table(rows=1, cols=2)
    table_edu.rows[0].cells[0].text = "Example University"
    table_edu.rows[0].cells[1].text = "2021 – 2026"

    doc.add_paragraph("EXPERIENCIA LABORAL")
    doc.add_paragraph("Automaticé los procesos de validación de facturas.")
    doc.add_paragraph("Extraje y analicé datos de SAP.")
    table_exp = doc.add_table(rows=1, cols=2)
    table_exp.rows[0].cells[0].text = "ExampleCorp — Intern"
    table_exp.rows[0].cells[1].text = "Nov 2024 – Feb 2025"

    path = tmp_path / "test_cv.docx"
    doc.save(str(path))
    return path


class TestReadCV:
    def test_parses_name(self, tmp_cv: Path):
        profile = read_cv(tmp_cv)
        assert profile.name == "ALEX SAMPLE CANDIDATE"

    def test_parses_contact(self, tmp_cv: Path):
        profile = read_cv(tmp_cv)
        assert "candidate@example.com" in profile.contact

    def test_parses_summary(self, tmp_cv: Path):
        profile = read_cv(tmp_cv)
        assert "Data Science" in profile.summary

    def test_detects_four_sections(self, tmp_cv: Path):
        profile = read_cv(tmp_cv)
        titles = [s.title for s in profile.sections]
        assert "EDUCACIÓN" in titles
        assert "EXPERIENCIA LABORAL" in titles

    def test_section_keeps_paragraphs(self, tmp_cv: Path):
        profile = read_cv(tmp_cv)
        exp = next(s for s in profile.sections if s.title == "EXPERIENCIA LABORAL")
        assert any("facturas" in p for p in exp.paragraphs)

    def test_section_keeps_tables(self, tmp_cv: Path):
        profile = read_cv(tmp_cv)
        exp = next(s for s in profile.sections if s.title == "EXPERIENCIA LABORAL")
        assert len(exp.tables) == 1
        assert exp.tables[0][0][0] == "ExampleCorp — Intern"

    def test_raw_text_contains_all_sections(self, tmp_cv: Path):
        profile = read_cv(tmp_cv)
        for title in ("EDUCACIÓN", "EXPERIENCIA LABORAL"):
            assert title in profile.raw_text

    def test_name_first_heading_not_treated_as_section(self, tmp_cv: Path):
        profile = read_cv(tmp_cv)
        titles = [s.title for s in profile.sections]
        assert "ALEX SAMPLE CANDIDATE" not in titles

    def test_to_json_round_trip(self, tmp_cv: Path, tmp_path: Path):
        import json

        profile = read_cv(tmp_cv)
        out = tmp_path / "profile.json"
        profile.to_json(out)
        # ensure the file is valid JSON
        loaded = json.loads(out.read_text(encoding="utf-8"))
        assert loaded["name"] == profile.name
        assert len(loaded["sections"]) == len(profile.sections)

    def test_cvprofile_dataclass_fields(self):
        p = CVProfile(name="x", contact="y", summary="z", sections=[], raw_text="t")
        assert p.name == "x"
        assert p.sections == []
        assert p.raw_text == "t"

    def test_cvsection_empty_detection(self):
        s = CVSection(title="X")
        assert s.is_empty() is True
        s.paragraphs.append("a")
        assert s.is_empty() is False


class TestIterBlockItemsOrder:
    def test_paragraphs_and_tables_interleave(self, tmp_cv: Path):
        from docx import Document

        doc = Document(str(tmp_cv))
        types = []
        for block in iter_block_items(doc):
            types.append(type(block).__name__)
        # Both kinds of block must appear in the iteration order
        assert "Paragraph" in types
        assert "Table" in types
        # The first table must come AFTER at least one paragraph — proving
        # python-docx default separate iteration (paragraphs-then-tables) is NOT
        # what we're getting; iter_block_items gives document order.
        first_table_idx = types.index("Table")
        assert "Paragraph" in types[:first_table_idx]
        # And a paragraph must follow the first table (otherwise tables would
        # all cluster at the end, which would defeat the purpose).
        assert "Paragraph" in types[first_table_idx + 1 :]