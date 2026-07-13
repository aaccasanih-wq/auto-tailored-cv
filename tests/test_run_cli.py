"""Tests for run.py — CLI plumbing and incremental cache logic.

We don't invoke the LLM or BrowserMCP here; we patch the subsystems.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import List
from unittest import mock

import pytest

import run as run_module
from run import (
    _is_processed,
    _job_cache_path,
    _load_cached_jobs,
    _save_job_cache,
    _tailor_one,
)
from src.config import settings
from src.extract.linkedin_scraper import SavedJob
from src.profile.cv_reader import CVProfile, CVSection
from tests.test_helpers_llm import StubLLMClient, llm_response


def _make_job(title="Senior Data Engineer", url="https://www.linkedin.com/jobs/view/12345/",
              company="Acme", job_id="12345") -> SavedJob:
    return SavedJob(
        title=title,
        url=url,
        company=company,
        location="Lima, Peru",
        description="Buscamos Data Engineer con SQL, Python, Snowflake.",
        job_id=job_id,
    )


@pytest.fixture
def jobs_dir(tmp_path):
    """Redirect settings.jobs_dir to a tmp_path for the duration of one test.
    Settings is frozen so we use object.__setattr__ and snapshot/restore
    manually.
    """
    fake = tmp_path / "jobs"
    fake.mkdir(parents=True, exist_ok=True)
    original = run_module.settings.jobs_dir
    object.__setattr__(run_module.settings, "jobs_dir", fake)
    try:
        yield fake
    finally:
        object.__setattr__(run_module.settings, "jobs_dir", original)


class TestCacheHelpers:
    def test_job_cache_path_uses_job_id(self):
        job = _make_job()
        p = _job_cache_path(job)
        assert p.name == "12345.json"
        assert p.parent == settings.jobs_dir

    def test_job_cache_path_falls_back_to_url(self):
        job = _make_job(job_id="", url="https://www.linkedin.com/jobs/view/abc-def/")
        p = _job_cache_path(job)
        assert p.name.endswith(".json")
        assert "abc" in p.name or "def" in p.name

    def test_unprocessed_when_no_cache(self, jobs_dir):
        assert _is_processed(_make_job()) is False

    def test_save_and_load_round_trip(self, jobs_dir):
        job = _make_job()
        _save_job_cache(job, tailored=False)
        loaded = _load_cached_jobs()
        assert len(loaded) == 1
        assert loaded[0].title == job.title
        assert loaded[0].job_id == job.job_id

    def test_is_processed_reads_tailored_flag(self, jobs_dir):
        _save_job_cache(_make_job(), tailored=True)
        assert _is_processed(_make_job()) is True

    def test_unmarked_job_is_not_processed(self, jobs_dir):
        _save_job_cache(_make_job(), tailored=False)
        assert _is_processed(_make_job()) is False


def _base_profile() -> CVProfile:
    return CVProfile(
        name="ALEX",
        contact="email",
        summary="x",
        sections=[
            CVSection(title="EDUCACIÓN", paragraphs=["a"], tables=[[["u", "y"]]]),
            CVSection(title="EXPERIENCIA LABORAL", paragraphs=["p"], tables=[[["c", "d"]]]),
            CVSection(title="HABILIDADES & HERRAMIENTAS", paragraphs=[], tables=[[["k1", "v1"]]]),
        ],
        raw_text="...",
    )


def _valid_tailored_json() -> str:
    return json.dumps({
        "summary": "x",
        "sections": [
            {"title": "EDUCACIÓN", "paragraphs": ["a"], "tables": [[["u", "y"]]]},
            {"title": "EXPERIENCIA LABORAL", "paragraphs": ["p"], "tables": [[["c", "d"]]]},
            {"title": "HABILIDADES & HERRAMIENTAS", "paragraphs": [], "tables": [[["k1", "v1"]]]},
        ],
    }, ensure_ascii=False)


@pytest.fixture
def output_dir(tmp_path):
    """Redirect settings.output_dir + base_cv_path for the duration of one test."""
    out = tmp_path / "output"
    out.mkdir(parents=True, exist_ok=True)
    base = _fake_base_docx(tmp_path)
    orig_out = run_module.settings.output_dir
    orig_base = run_module.settings.base_cv_path
    object.__setattr__(run_module.settings, "output_dir", out)
    object.__setattr__(run_module.settings, "base_cv_path", base)
    try:
        yield out
    finally:
        object.__setattr__(run_module.settings, "output_dir", orig_out)
        object.__setattr__(run_module.settings, "base_cv_path", orig_base)


class TestTailorOne:
    def test_dry_run_creates_job_description_only(self, output_dir):
        job = _make_job()
        out = _tailor_one(client=None, base_profile=_base_profile(), job=job, dry_run=True)
        assert out is not None
        assert (out / "job_description.txt").exists()
        # No analysis.json in dry-run
        assert not (out / "analysis.json").exists()

    def test_full_run_produces_docx(self, output_dir):
        # Stub the LLM: tailor + evaluator (pass) — no repair needed.
        stub = StubLLMClient([
            llm_response(_valid_tailored_json()),
            llm_response(json.dumps({"issues": [], "overall_verdict": "pass", "summary": "ok"})),
        ])
        job = _make_job()
        out = _tailor_one(client=stub, base_profile=_base_profile(), job=job, dry_run=False)
        assert out is not None
        assert (out / "cv.docx").exists()
        assert (out / "analysis.json").exists()
        assert (out / "evaluation.json").exists()
        # No repaired file when verdict is pass
        assert not (out / "analysis_repaired.json").exists()

    def test_repair_pass_writes_repaired_json(self, output_dir, monkeypatch):
        stub = StubLLMClient([
            llm_response(_valid_tailored_json()),
            llm_response(json.dumps({
                "issues": [{
                    "id": "1", "type": "verbatim_copy", "severity": "high",
                    "quote": "pipelines ETL", "explanation": "copied",
                    "suggested_fix": "paraphrase"
                }],
                "overall_verdict": "needs_repair", "summary": "x"
            })),
            llm_response(_valid_tailored_json()),  # repair returns corrected
        ])
        monkeypatch.setattr(run_module, "make_client", lambda: stub)
        out = _tailor_one(client=stub, base_profile=_base_profile(),
                          job=_make_job(), dry_run=False)
        assert (out / "analysis_repaired.json").exists()

    def test_tailored_with_no_sections_is_failure(self, output_dir):
        bad = json.dumps({"summary": "x", "sections": []})
        stub = StubLLMClient([llm_response(bad)])
        result = _tailor_one(client=stub, base_profile=_base_profile(),
                              job=_make_job(), dry_run=False)
        assert result is None


def _fake_base_docx(tmp_path: Path) -> Path:
    """Build a minimal docx fixture for renderer tests."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("ALEX")
    doc.add_paragraph("email")
    doc.add_paragraph("x")
    doc.add_paragraph("EDUCACIÓN")
    doc.add_paragraph("a")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "u"
    t.rows[0].cells[1].text = "y"
    doc.add_paragraph("EXPERIENCIA LABORAL")
    doc.add_paragraph("p")
    t2 = doc.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "c"
    t2.rows[0].cells[1].text = "d"
    doc.add_paragraph("HABILIDADES & HERRAMIENTAS")
    t3 = doc.add_table(rows=1, cols=2)
    t3.rows[0].cells[0].text = "k1"
    t3.rows[0].cells[1].text = "v1"
    path = tmp_path / "base_cv.docx"
    doc.save(str(path))
    return path


class TestArgparseAndCommands:
    def test_parser_requires_subcommand(self):
        with pytest.raises(SystemExit):
            run_module.build_parser().parse_args([])

    def test_extract_subcommand(self):
        p = run_module.build_parser()
        args = p.parse_args(["extract"])
        assert args.cmd == "extract"

    def test_all_with_flags(self):
        p = run_module.build_parser()
        args = p.parse_args(["all", "--new", "--force", "--dry-run",
                             "--job", "https://example.com", "--limit", "3"])
        assert args.cmd == "all"
        assert args.new is True
        assert args.force is True
        assert args.dry_run is True
        assert args.job == "https://example.com"
        assert args.limit == 3

    def test_tailor_only_command(self):
        p = run_module.build_parser()
        args = p.parse_args(["tailor"])
        assert args.cmd == "tailor"
        assert args.new is False

    def test_main_smoke_for_tailor_no_jobs(self, tmp_path):
        # Patch settings.jobs_dir to empty so _load_cached_jobs returns []
        from run import main
        original = run_module.settings.jobs_dir
        object.__setattr__(run_module.settings, "jobs_dir", tmp_path / "nojobs")
        try:
            rc = main(["tailor"])
            assert rc == 0  # no jobs → 0, just a warning
        finally:
            object.__setattr__(run_module.settings, "jobs_dir", original)