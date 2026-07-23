"""Build a simplified Harvard-style base CV from scratch.

The original CV had several formatting quirks that were causing the docx
renderer to break, especially when paragraphs had:
  - multiple runs with mixed fonts/sizes (one emptied run carried stray style)
  - tab stops with leading spaces for alignment
  - multi-line cells (literal \n inside a cell's text)
  - justify alignment that LibreOffice renders with big intraword gaps

This new base_cv.docx uses a single Calibri 10.5pt baseline, one run per
paragraph, left-aligned body text, and single-line cells. All the FACTS
from the original CV are preserved.

Run once. Output: input/base_cv.docx (overwrites).
The original is preserved at input/base_cv_original_backup.docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BODY_FONT = "Calibri"
BODY_SIZE = Pt(10.5)
HEADING_SIZE = Pt(11)
NAME_SIZE = Pt(16)
LINE_BREAK_REPLACEMENT = " · "


def _set_run(run, *, bold=False, italic=False, font=BODY_FONT, size=BODY_SIZE):
    run.font.name = font
    run.font.size = size
    run.bold = bold
    run.italic = italic
    # Force east asian font too so LibreOffice doesn't substitute
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run(run)
    return p


def _add_heading(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    _set_run(run, bold=True, size=HEADING_SIZE)
    return p


def _add_simple_paragraph(doc, text, *, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=BODY_SIZE):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    _set_run(run, bold=bold, italic=italic, size=size)
    return p


def _add_simple_table(doc, rows, col_widths=None):
    """rows: list of list of str. Borders invisible."""
    n_rows = len(rows)
    n_cols = len(rows[0]) if rows else 0
    if n_rows == 0 or n_cols == 0:
        return None
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"  # we'll hide borders next
    table.autofit = True
    # Remove all borders
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        # Set tblBorders to "none" everywhere
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
            tblPr.append(tblBorders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge_el = tblBorders.find(qn(f"w:{edge}"))
            if edge_el is None:
                edge_el = tblBorders.makeelement(qn(f"w:{edge}"), {})
                tblBorders.append(edge_el)
            edge_el.set(qn("w:val"), "none")
            edge_el.set(qn("w:sz"), "0")
            edge_el.set(qn("w:space"), "0")
            edge_el.set(qn("w:color"), "auto")
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            # Write text in the first paragraph
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Strip any existing content
            for r in list(cell.paragraphs[0].runs):
                r.text = ""
            run = cell.paragraphs[0].add_run(cell_text)
            _set_run(run, bold=(ci == 0 and ri == 0) if False else False)
            # First column of every row bold (it's the category/role label)
            if ci == 0:
                _set_run(run, bold=True)
            else:
                _set_run(run, bold=False)
    return table


def build():
    out_path = Path(__file__).resolve().parent.parent / "input" / "base_cv.docx"
    doc = Document()

    # Set default Normal style font
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE

    # Tighten default paragraph spacing slightly for a Harvard compact look
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15

    # --- HEADER (name + contact + summary, all centered) ---
    _add_simple_paragraph(doc, "AXEL AARON CCASANI HUACHUA",
                          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=NAME_SIZE)
    _add_simple_paragraph(doc, "986 531 450 | aaronaxel810@gmail.com | Lima, Peru",
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_simple_paragraph(doc, "En búsqueda de un puesto en Digital Products · Análisis de Datos · Transformación Digital",
                          italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Spacer
    doc.add_paragraph("")

    # --- EDUCACIÓN ---
    _add_heading(doc, "EDUCACIÓN")
    _add_simple_table(doc, [
        ["Universidad del Pacífico — Lima, Peru", "2021 – 2026"],
    ])
    _add_simple_paragraph(
        doc,
        "Lic. en Economía  |  En proceso de obtener el grado de bachiller (Julio 2026)",
        italic=True,
    )

    doc.add_paragraph("")

    # --- EXPERIENCIA LABORAL ---
    _add_heading(doc, "EXPERIENCIA LABORAL")
    # Metso
    _add_simple_table(doc, [["Metso Perú — Practicante de Cuentas por Pagar", "Nov 2024 – Feb 2025"]])
    _add_bullet(doc, "Automaticé los procesos de validación de facturas (basados en Excel), reduciendo errores manuales y tiempos de procesamiento en 50%.")
    _add_bullet(doc, "Extraje y analicé datos de SAP para generar informes financieros, mejorando la visibilidad del estado de las cuentas por pagar.")
    _add_bullet(doc, "Gestioné y di seguimiento integral a facturas, resolviendo discrepancies y garantizando pagos puntuales.")
    # Nissan
    _add_simple_table(doc, [["Nissan Perú — Practicante de Finanzas", "Dic 2023 – Abr 2024"]])
    _add_bullet(doc, "Detecté pagos incorrectos de bonificaciones y descuentos (~20% del precio base del vehículo), evitando pérdidas significativas de ingresos.")
    _add_bullet(doc, "Implementé una herramienta en Excel para validar pagos a concesionarios (bonificaciones/descuentos).")
    _add_bullet(doc, "Realicé análisis de rentabilidad, identificando los factores que influyeron en la brecha entre las ventas proyectadas y las reales.")
    _add_bullet(doc, "Optimicé procesos internos, reduciendo el tiempo de validación de pagos (bonificaciones/descuentos) de más de 7 días a solo 2 días, mejorando la eficiencia operativa.")

    doc.add_paragraph("")

    # --- PROYECTOS ---
    _add_heading(doc, "PROYECTOS")

    # Project 1: Rastreador
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run("Rastreador de Gastos Automatizado con IA")
    _set_run(r1, bold=True)
    r2 = p.add_run(" (Dashboard) · May 2026")
    _set_run(r2)
    _add_bullet(doc, "Desarrollé una automatización que analiza correos de notificación bancaria usando la API de Gmail y registra transacciones de gastos en Google Sheets mediante DeepSeek AI para extracción de lenguaje natural. Además, integra un bot de Telegram que acepta mensajes de texto o audio para registrar gastos manualmente en caso de ser necesario (gastos en efectivo).")
    _add_bullet(doc, "Desarrollé un dashboard interactivo en Streamlit para visualización y análisis del gasto en tiempo real, permitiendo tomar decisiones financieras personales basadas en datos por categorías de gasto.")

    # Project 2: KAYLA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run("KAYLA — Recordatorios automáticos de citas y medicamentos")
    _set_run(r1, bold=True)
    r2 = p.add_run(" (Landing Page · Dashboard) · Jun 2026")
    _set_run(r2)
    _add_bullet(doc, "Diseñé un sistema que automatiza recordatorios de citas y recojo de medicamentos para postas de salud: un Google Form alimenta la base de pacientes en Google Sheets, y un scheduler en GitHub Actions detecta a quienes tienen fecha próxima y envía recordatorios automáticos vía bot de Telegram al médico a cargo.")
    _add_bullet(doc, "Construí un dashboard en Streamlit que consume la base de pacientes y monitorea en tiempo real el estado de cada recordatorio y paciente, equivalente al seguimiento de métricas de un funnel de conversión.")

    # Project 3: AI Personal Agent
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run("AI Personal Agent")
    _set_run(r1, bold=True)
    r2 = p.add_run(" (Agentic AI · RAG · Automatización) · Jun 2026")
    _set_run(r2)
    _add_bullet(doc, "Desarrollé un asistente de IA modular basado en DeepSeek que procesa entradas de voz y texto para gestionar correos electrónicos, archivos locales, recordatorios diarios y más. Casi todas las funciones se recuperan dinámicamente mediante RAG según la intención del usuario.")
    _add_bullet(doc, "Diseñé una arquitectura modular tipo LEGO, donde cada herramienta es un script independiente. Implementa la recuperación de memoria de conversaciones pasadas basada en RAG, con un marco extensible para la generación autónoma de herramientas.")

    doc.add_paragraph("")

    # --- HABILIDADES & HERRAMIENTAS ---
    _add_heading(doc, "HABILIDADES & HERRAMIENTAS")
    # One-row-per-category table. Cells are SINGLE LINE (no embedded \n).
    _add_simple_table(doc, [
        ["Programming", "Python (Pandas, NumPy, Matplotlib, Seaborn, Selenium, APIs/Requests, GeoPandas, Folium, Raster, Langchain, etc), SQL, R (RStudio), Excel"],
        ["Data Science", "Web scraping, API, data cleaning & transformation, statistical analysis, data visualization, geospatial analysis (static & dynamic maps), etc"],
        ["Finance & BI", "Financial statement analysis, project valuation, profitability analysis, dashboard & report design"],
        ["Software", "SAP, Bloomberg Terminal, GitHub, Claude Desktop, Opencode"],
        ["Inteligencia Artificial", "Prompting con LLMs, automatización con APIs, RAG"],
        ["Idiomas", "Español (nativo), Inglés (Avanzado)"],
    ])

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    build()