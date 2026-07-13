"""Render a tailored CV JSON back into a .docx, preserving the base template's
formatting (fonts, sizes, margins, alignment) by replacing text run-by-run.

Strategy:
  1. Open base_cv.docx in-memory (the existing file acts as our template).
  2. Walk body blocks (paragraphs + tables) in document order.
  3. The first 3 non-heading paragraphs before any section heading are
     name / contact / summary — name and contact are left UNTOUCHED because
     they are immutable personal facts; the summary line is overwritten with
     the tailored summary. Any extra pre-header paragraphs beyond summary
     are cleared.
  4. Section headings (all-caps paragraphs) are left UNTOUCHED — they mark
     sections AND are already correct titles.
  5. Each content paragraph inside a section is replaced with the matching
     `tailored["sections"][i]["paragraphs"][j]` string. Extras are blanked.
  6. Each content table inside a section is replaced cell-by-cell using the
     matching `tailored["sections"][i]["tables"][k]` grid, preserving the
     per-cell run formatting from the first paragraph / first run.

Cell text replacement is a single-paragraph collapse: we strip extra
paragraphs inside a cell and write the full new text into the first
paragraph's first run, inheriting its formatting.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from src.profile.cv_reader import iter_block_items, _is_heading
from src.utils.logging import get_logger

log = get_logger(__name__)


# --------------------------------------------------------------------------- #
# Low-level text replacement                                                   #
# --------------------------------------------------------------------------- #


def _set_paragraph_text(p: Paragraph, new_text: str) -> None:
    """Replace a paragraph's text while keeping the formatting of its first run."""
    runs = list(p.runs)
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _set_cell_text(cell: _Cell, new_text: str) -> None:
    """Replace a cell's content with new_text. We keep the first paragraph and
    its first run's formatting, then DELETE any extra paragraphs from the cell
    XML. Cells with extra (now-empty) paragraphs would otherwise leak into
    cell.text as trailing newlines."""
    paras = list(cell.paragraphs)
    if not paras:
        return
    first = paras[0]
    _set_paragraph_text(first, new_text)
    # Drop extra paragraphs entirely so cell.text returns just our single line.
    for p in paras[1:]:
        p._element.getparent().remove(p._element)


# --------------------------------------------------------------------------- #
# High-level writer                                                           #
# --------------------------------------------------------------------------- #


def write_tailored_docx(
    base_cv_path: Path,
    tailored_json: Dict[str, Any],
    output_path: Path,
) -> Path:
    """Produce a tailored .docx at `output_path` based on `base_cv_path`.

    The base file is opened in-memory via python-docx and modified block-by-block.
    The original file on disk is preserved.
    """
    if not base_cv_path.exists():
        raise FileNotFoundError(f"base CV not found at {base_cv_path}")
    doc = Document(str(base_cv_path))

    summary = tailored_json.get("summary", "") or ""
    sections = tailored_json.get("sections", []) or []

    section_idx = -1
    name_assigned = False
    para_idx_in_section = 0
    table_idx_in_section = 0
    pre_header_para_count = 0
    used_sections = 0

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if _is_heading(text):
                # The FIRST all-caps line is the person's name (matches the
                # reader's convention). After the name, all-caps lines are
                # actual section headings.
                if not name_assigned:
                    name_assigned = True
                    continue
                # Start a new section. Whatever used_sections comes from the
                # tailored_json side is the index of the section we'll feed.
                section_idx = used_sections
                used_sections += 1
                para_idx_in_section = 0
                table_idx_in_section = 0
                # Headings are immutable.
                continue
            if section_idx < 0:
                # Pre-header layout (matches cv_reader's convention):
                #   name (already consumed as heading), then contact, then summary,
                #   then any extra pre-header paragraphs which we blank out.
                pre_header_para_count += 1
                if pre_header_para_count == 1:
                    # contact: do NOT modify
                    continue
                if pre_header_para_count == 2:
                    _set_paragraph_text(block, summary)
                else:
                    # Extra pre-header paragraphs (spacers, etc.). Clear them so
                    # the read-back summary isn't bloated with leftover text.
                    _set_paragraph_text(block, "")
            else:
                section_data: Optional[Dict[str, Any]] = (
                    sections[section_idx] if section_idx < len(sections) else None
                )
                paragraphs = (section_data or {}).get("paragraphs", []) or []
                if para_idx_in_section < len(paragraphs):
                    _set_paragraph_text(block, paragraphs[para_idx_in_section])
                else:
                    # extras that the LLM should have left empty
                    _set_paragraph_text(block, "")
                para_idx_in_section += 1
        elif isinstance(block, Table):
            if section_idx < 0:
                # Pre-header tables are unusual; leave untouched.
                continue
            section_data = sections[section_idx] if section_idx < len(sections) else None
            tables = (section_data or {}).get("tables", []) or []
            if table_idx_in_section >= len(tables):
                # No tailored grid for this table — leave untouched.
                table_idx_in_section += 1
                continue
            tailored_grid = tables[table_idx_in_section]
            _replace_table(block, tailored_grid)
            table_idx_in_section += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("docx written: %s", output_path)
    return output_path


def _replace_table(table: Table, grid: List[List[str]]) -> None:
    """Replace each cell in `table` with the text in `grid`, by position."""
    rows = list(table.rows)
    if len(rows) != len(grid):
        log.warning(
            "table shape mismatch (rows): docx=%d tailored=%d — replacing only the min",
            len(rows),
            len(grid),
        )
    for r_idx, row in enumerate(rows):
        if r_idx >= len(grid):
            # leave remaining rows as-is
            break
        tailored_row = grid[r_idx] or []
        cells = list(row.cells)
        if len(cells) != len(tailored_row):
            log.debug(
                "row %d col mismatch: docx=%d tailored=%d — replacing min",
                r_idx,
                len(cells),
                len(tailored_row),
            )
        for c_idx, cell in enumerate(cells):
            if c_idx >= len(tailored_row):
                break
            _set_cell_text(cell, tailored_row[c_idx])


__all__ = ["write_tailored_docx"]