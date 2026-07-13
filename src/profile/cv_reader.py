"""Read and structure a base CV in .docx format.

The base CV is read with python-docx. Paragraphs and tables are walked in
document order (not separately) so the resulting structured representation
preserves the original layout. Section headings are detected heuristically:
ALL-CAPS Normal paragraphs are treated as headings. This matches Axel's CV
convention but also works for most résumé templates.

Output: a CVProfile with a flat list of CVSection objects, plus a plain-text
rendering used for the LLM prompt context.
"""

from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Iterator, List, Optional, Union

from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


_LETTER_RE = re.compile(r"[A-Z]")
_LOWER_RE = re.compile(r"[a-zà-ÿ]")  # ascii lowercase + accented lowercase (spanish)


@dataclass
class CVSection:
    """A detected section of the CV: heading + its content blocks."""
    title: str
    paragraphs: List[str] = field(default_factory=list)
    tables: List[List[List[str]]] = field(default_factory=list)

    def is_empty(self) -> bool:
        return not self.paragraphs and not self.tables

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class CVProfile:
    """Structured representation of a base CV."""
    name: str
    contact: str
    summary: str
    sections: List[CVSection]
    raw_text: str

    def to_dict(self) -> dict:
        return {
            "name": self.name,
            "contact": self.contact,
            "summary": self.summary,
            "sections": [s.to_dict() for s in self.sections],
            "raw_text": self.raw_text,
        }

    def to_json(self, path: Optional[Path] = None) -> str:
        data = self.to_dict()
        if path:
            path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return json.dumps(data, ensure_ascii=False, indent=2)


def iter_block_items(parent: Union[_Document, _Cell]) -> Iterator[Union[Paragraph, Table]]:
    """Yield Paragraph and Table objects in document order."""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("unsupported parent type")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _is_heading(text: str) -> bool:
    """A heading is a non-empty line that contains uppercase letters but
    NO lowercase letters (case-insensitive digits/punctuation allowed).
    Handles Spanish accents (à-ÿ).
    """
    stripped = text.strip()
    if not stripped:
        return False
    if not _LETTER_RE.search(stripped):
        return False  # no letters → not a heading
    return not _LOWER_RE.search(stripped)


def _table_to_grid(table: Table) -> List[List[str]]:
    grid: List[List[str]] = []
    for row in table.rows:
        grid.append([cell.text.strip() for cell in row.cells])
    return grid


def read_cv(path: Path) -> CVProfile:
    """Read a .docx CV and return a structured CVProfile."""
    from docx import Document

    doc = Document(str(path))

    name = ""
    contact = ""
    summary = ""
    sections: List[CVSection] = []
    current: Optional[CVSection] = None
    name_assigned = False
    # Non-heading lines that appear before any section heading form the header
    pre_header_lines: List[str] = []

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if _is_heading(text):
                # First all-caps line is the person's name (common CV convention).
                # Subsequent all-caps lines are section headings.
                if not name_assigned:
                    name = text
                    name_assigned = True
                    continue
                current = CVSection(title=text)
                sections.append(current)
                continue
            if current is not None:
                if text:
                    current.paragraphs.append(text)
            else:
                if text:
                    pre_header_lines.append(text)
        elif isinstance(block, Table):
            grid = _table_to_grid(block)
            if not grid:
                continue
            if current is not None:
                current.tables.append(grid)
            else:
                # Tables before any heading — unusual for a CV; treat as header context
                # by collapsing them into pre_header_lines
                for row in grid:
                    pre_header_lines.append(" | ".join(cell for cell in row if cell))

    # After the first heading (which we repurposed as the name), the lines that
    # appeared before any further section heading are the contact + summary.
    if len(pre_header_lines) >= 1:
        contact = pre_header_lines[0]
    if len(pre_header_lines) >= 2:
        summary = "\n".join(pre_header_lines[1:])

    # Drop trailing empty sections just in case
    sections = [s for s in sections if not s.is_empty()] + [
        s for s in sections if s.is_empty()
    ]

    raw_text = _render_text(name, contact, summary, sections)

    return CVProfile(
        name=name,
        contact=contact,
        summary=summary,
        sections=sections,
        raw_text=raw_text,
    )


def _render_text(name: str, contact: str, summary: str, sections: List[CVSection]) -> str:
    """Build a clean plain-text representation of the CV."""
    parts: List[str] = []
    if name:
        parts.append(name)
    if contact:
        parts.append(contact)
    if summary:
        parts.append("")
        parts.append(summary)
    for section in sections:
        parts.append("")
        parts.append(section.title)
        parts.append("=" * len(section.title))
        for p in section.paragraphs:
            parts.append(f"- {p}")
        for table in section.tables:
            for row in table:
                parts.append(" | ".join(cell for cell in row if cell))
    return "\n".join(parts).strip() + "\n"


__all__ = ["CVProfile", "CVSection", "read_cv", "iter_block_items"]